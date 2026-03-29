#!/usr/bin/env python3
from __future__ import annotations

"""
Aircraft Checklist Converter
=============================
Converts aircraft checklist files (PDF, Excel, Google Sheets, Word, TXT)
into structured CSV or JSON compatible with FlightCheck.

Usage:
    python checklist_converter.py                     # Interactive mode
    python checklist_converter.py path/to/file.pdf    # Direct file
    python checklist_converter.py <google-sheets-url> # Google Sheets

Output CSV format matches FlightCheck's import format:
    name, manufacturer, type, image, checklist category, phase, item, expectedState, notes
"""

import argparse
import csv
import io
import json
import os
import re
import sys
from pathlib import Path

# Ensure Unicode output works on Windows terminals
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

# ---------------------------------------------------------------------------
# Dependency checks — give clear install instructions if missing
# ---------------------------------------------------------------------------

MISSING_DEPS: list[str] = []


def _check_dep(module: str, pip_name: str) -> bool:
    try:
        __import__(module)
        return True
    except ImportError:
        MISSING_DEPS.append(pip_name)
        return False


HAS_PDFPLUMBER = _check_dep("pdfplumber", "pdfplumber")
HAS_DOCX = _check_dep("docx", "python-docx")
HAS_OPENPYXL = _check_dep("openpyxl", "openpyxl")
HAS_REQUESTS = _check_dep("requests", "requests")
HAS_FITZ = _check_dep("fitz", "pymupdf")

# ---------------------------------------------------------------------------
# Data model (mirrors FlightCheck's TypeScript types)
# ---------------------------------------------------------------------------


class ChecklistItem:
    def __init__(self, label: str, expected_state: str = "", notes: str = ""):
        self.label = label
        self.expected_state = expected_state
        self.notes = notes

    def to_dict(self) -> dict:
        d: dict = {"label": self.label}
        if self.expected_state:
            d["expectedState"] = self.expected_state
        if self.notes:
            d["notes"] = self.notes
        return d


class ChecklistPhase:
    def __init__(self, title: str, category: str = ""):
        self.title = title
        self.category = category
        self.items: list[ChecklistItem] = []

    def to_dict(self) -> dict:
        phase_id = re.sub(r"[^a-z0-9]+", "-", self.title.lower()).strip("-")
        return {
            "id": phase_id,
            "title": self.title,
            "items": [
                {
                    "id": f"{phase_id}-{i}",
                    **item.to_dict(),
                }
                for i, item in enumerate(self.items)
            ],
        }


class ParsedChecklist:
    def __init__(
        self,
        plane_name: str = "",
        manufacturer: str = "",
        plane_type: str = "GA",
        category: str = "",
    ):
        self.plane_name = plane_name
        self.manufacturer = manufacturer
        # When True (default), phases with no category are skipped if other
        # phases have a category — used to drop pre-header noise from PDFs.
        # Set to False for structured parsers (HTML, Excel, CSV) where empty
        # category means "Normal" and all phases should be included.
        self.skip_uncategorized: bool = True
        self.plane_type = plane_type
        self.category = category
        self.phases: list[ChecklistPhase] = []

    @property
    def total_items(self) -> int:
        return sum(len(p.items) for p in self.phases)

    def to_flightcheck_json(self) -> dict:
        plane_id = re.sub(r"[^a-z0-9]+", "-", self.plane_name.lower()).strip("-")
        return {
            "plane": {
                "id": plane_id,
                "name": self.plane_name,
                "manufacturer": self.manufacturer,
                "image": "",
                "type": self.plane_type,
            },
            "checklist": {
                "planeId": plane_id,
                "phases": [p.to_dict() for p in self.phases],
            },
        }

    def to_csv_rows(self) -> list[list[str]]:
        rows: list[list[str]] = []
        # When the document has explicit section headers (Normal/Abnormal/Emergency
        # Checklist), optionally drop phases that appeared before the first header —
        # these are typically reference data or noise from PDF/text extraction.
        # Structured parsers (HTML, Excel, CSV) set skip_uncategorized=False so all
        # phases are included regardless of whether a category is set.
        has_sections = any(p.category for p in self.phases)
        for phase in self.phases:
            cat = phase.category if phase.category else self.category
            if self.skip_uncategorized and has_sections and not cat:
                continue
            for item in phase.items:
                rows.append(
                    [
                        self.plane_name,
                        self.manufacturer,
                        self.plane_type,
                        "",  # image
                        cat,
                        item.label,
                        item.expected_state,
                        item.notes,
                        phase.title,
                    ]
                )
        return rows


# ---------------------------------------------------------------------------
# Text extraction by file type
# ---------------------------------------------------------------------------


def extract_pdf_text(filepath: str) -> str:
    if not HAS_PDFPLUMBER:
        raise RuntimeError("pdfplumber is required for PDF files: pip install pdfplumber")
    import pdfplumber

    lines: list[str] = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            # Prefer text extraction: handles pages with mixed table+checklist content
            # correctly (table extraction drops non-table text on the same page).
            text = page.extract_text()
            if text and text.strip():
                lines.append(text)
            else:
                # Fall back to table extraction for pages with no extractable text
                for table in page.extract_tables():
                    for row in table:
                        cells = [c.strip() if c else "" for c in row]
                        if any(cells):
                            lines.append("\t".join(cells))
    return "\n".join(lines)


def extract_docx_text(filepath: str) -> str:
    if not HAS_DOCX:
        raise RuntimeError("python-docx is required for Word files: pip install python-docx")
    from docx import Document

    doc = Document(filepath)
    lines: list[str] = []

    # Extract from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Extract from tables (common in structured checklists)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("\t".join(cells))

    return "\n".join(lines)


def extract_excel_data(filepath: str) -> tuple[list[str], list[list[str]]]:
    """Returns (headers, rows) from the first sheet with data."""
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl is required for Excel files: pip install openpyxl")
    import openpyxl

    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    ws = wb.active
    if ws is None:
        raise RuntimeError("Excel file has no active sheet")

    all_rows: list[list[str]] = []
    for row in ws.iter_rows(values_only=True):
        cells = [str(c).strip() if c is not None else "" for c in row]
        if any(cells):
            all_rows.append(cells)
    wb.close()

    if not all_rows:
        raise RuntimeError("Excel file is empty")

    return all_rows[0], all_rows[1:]


def download_google_sheet(url: str) -> str:
    """Download a Google Sheet as CSV using its public export URL."""
    if not HAS_REQUESTS:
        raise RuntimeError("requests is required for Google Sheets: pip install requests")
    import requests

    # Extract the sheet ID from various URL formats
    patterns = [
        r"/spreadsheets/d/([a-zA-Z0-9_-]+)",
        r"id=([a-zA-Z0-9_-]+)",
    ]
    sheet_id = None
    for pat in patterns:
        m = re.search(pat, url)
        if m:
            sheet_id = m.group(1)
            break

    if not sheet_id:
        raise ValueError(
            "Could not extract Google Sheet ID from URL.\n"
            "Expected format: https://docs.google.com/spreadsheets/d/<SHEET_ID>/..."
        )

    # Extract gid (sheet tab) if present, default to 0
    gid_match = re.search(r"gid=(\d+)", url)
    gid = gid_match.group(1) if gid_match else "0"

    export_url = (
        f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}"
    )

    print(f"  Downloading from Google Sheets (ID: {sheet_id}, gid: {gid})...")
    resp = requests.get(export_url, timeout=30)

    if resp.status_code == 200 and "text/" in resp.headers.get("content-type", ""):
        return resp.text
    elif resp.status_code == 200:
        # Might be an HTML login page — sheet is not public
        raise RuntimeError(
            "Google Sheet is not publicly accessible.\n"
            "Go to Share > 'Anyone with the link' > Viewer, then try again.\n"
            "Or export it manually as CSV/Excel and pass the file path instead."
        )
    else:
        raise RuntimeError(
            f"Failed to download Google Sheet (HTTP {resp.status_code}).\n"
            "Make sure the sheet is shared as 'Anyone with the link'."
        )


# ---------------------------------------------------------------------------
# Heuristic text parser (ported from FlightCheck's checklistFileParser.ts)
# ---------------------------------------------------------------------------

PHASE_KEYWORDS = re.compile(
    r"^(pre[- ]?flight|before\s|after\s|engine\s|taxi|takeoff|take-off|"
    r"landing|climb|cruise|descent|approach|shutdown|securing|startup|"
    r"run[- ]?up|normal|emergency|abnormal|ground\s|parking|ramp|"
    r"starting|before\s+start|after\s+start|in[- ]?flight|go[- ]?around|"
    r"missed\s+approach|holding|balked\s+landing|"
    r"short\s|enroute\s|air\s+start|asymmetric|precautionary|ditching|"
    r"generator|inverter|starter\s|battery\s|cabin\s|wing\s+fire|"
    r"electrical\s+fire|fcu\s|loss\s+of|fuel\s+reservoir|"
    r"upper\s+half|lower\s+half|crew\s+door|cargo\s+pod|"
    r"windmill|overheated|flaps\s+fail)",
    re.IGNORECASE,
)

# Detects top-level checklist section headers (become the CSV `category` field)
_SECTION_RE = re.compile(
    r"\b(normal|abnormal|emergency)\s+checklist\b",
    re.IGNORECASE,
)

# Phase headers that use "Title: Sub-condition" format (e.g., "Emergency descent: Rough air")
# All other "Label: Value" lines are treated as items, not phase headers
_COLON_PHASE_RE = re.compile(
    r"^(emergency\s+(descent|landing)|engine\s+(start|fire)|"
    r"air\s+start|enroute\s+climb|short\s+(takeoff|take-off|field))\s*:",
    re.IGNORECASE,
)


def _is_phase_header(line: str) -> bool:
    # Conditional lines ("If X:") are never phase headers
    if re.match(r"^[Ii]f\s", line):
        return False

    # Lines with "Label: Value" format are items, not headers — with one exception:
    # specific procedure titles like "Emergency descent: Rough air"
    if re.search(r":\s+\S", line):
        return bool(_COLON_PHASE_RE.match(line))

    clean = line.rstrip(":").strip()

    # ALL CAPS with at least 3 letter characters
    letters = re.sub(r"[^a-zA-Z]", "", clean)
    if len(letters) >= 3 and clean == clean.upper() and re.search(r"[A-Z]", clean):
        # Not an item (no dot/dash/tab separators, or challenge-response dash)
        if not re.search(r"\.{3,}|_{3,}|-{3,}|\t|\s+-\s+", line):
            return True

    # Ends with colon and is short (likely a header)
    if line.endswith(":") and len(line) < 60 and "..." not in line:
        return True

    # Matches common phase keywords (but not if it looks like a challenge-response item or tab-row)
    if PHASE_KEYWORDS.search(clean) and len(clean) < 50 and not re.search(r"\.{3,}|_{3,}|\t|\s+-\s+", line):
        return True

    return False


def _clean_phase_title(line: str) -> str:
    title = line.rstrip(":").strip()
    title = re.sub(r"^\d+[.)]\s*", "", title)  # Remove leading numbers
    # Title case
    return title.title()


def _parse_item_line(line: str) -> tuple[str, str] | None:
    if len(line) < 5:
        return None

    # Pattern 1: "Label ......... STATE" or "Label _____ STATE"
    m = re.match(r"^(.+?)\s*[._]{3,}\s*(.+)$", line)
    if m:
        return m.group(1).strip(), m.group(2).strip()

    # Pattern 1b: "Label: Value" (colon + space) — common in PDF checklists
    m = re.match(r"^(.{2,}?):\s+(.+)$", line)
    if m:
        return m.group(1).strip(), m.group(2).strip()

    # Pattern 2: "Label - STATE" or "Label -- STATE"
    m = re.match(r"^(.{2,}?)\s+-{1,3}\s+([A-Z][A-Z0-9 /()\-&_,.]{1,40})$", line)
    if m:
        return m.group(1).strip(), m.group(2).strip()

    # Pattern 3: Tab-separated "Label\tSTATE" (pdfplumber table extraction)
    m = re.match(r"^(.+?)\t+(.+)$", line)
    if m and len(m.group(2).strip()) < 40:
        return m.group(1).strip().rstrip(":"), m.group(2).strip()

    # Pattern 4: "Label    STATE" (3+ spaces, state is ALL CAPS)
    m = re.match(r"^(.{5,}?)\s{3,}([A-Z][A-Z0-9 /()\-&_,.]{1,40})$", line)
    if m:
        return m.group(1).strip(), m.group(2).strip()

    # Pattern 5: Bullet/checkbox items
    m = re.match(r"^[-\u2022\u25cb\u25a1\u2610\u25a0\u25cf]\s+(.+)$", line)
    if m:
        return m.group(1).strip(), ""

    # Pattern 6: Numbered items "1. Label" or "1) Label"
    m = re.match(r"^\d+[.)]\s+(.+)$", line)
    if m and not _is_phase_header(line):
        return m.group(1).strip(), ""

    return None


def parse_text_to_checklist(text: str) -> list[ChecklistPhase]:
    """Heuristic parser: extract phases and items from raw text."""
    raw_lines = [l.strip() for l in text.split("\n") if l.strip()]
    phases: list[ChecklistPhase] = []
    current_phase: ChecklistPhase | None = None
    current_category: str = ""

    for line in raw_lines:
        if len(line) < 3:
            continue

        # Detect top-level section headers (Normal/Abnormal/Emergency Checklist)
        section_match = _SECTION_RE.search(line)
        if section_match:
            kind = section_match.group(1).title()
            current_category = f"{kind} Checklist"
            continue

        if _is_phase_header(line):
            title = _clean_phase_title(line)
            if len(title) < 2:
                continue
            current_phase = ChecklistPhase(title, category=current_category)
            phases.append(current_phase)
            continue

        parsed = _parse_item_line(line)
        if parsed:
            if current_phase is None:
                current_phase = ChecklistPhase("General", category=current_category)
                phases.append(current_phase)
            label, state = parsed
            current_phase.items.append(ChecklistItem(label, state))

    # Filter empty phases
    return [p for p in phases if p.items]


# ---------------------------------------------------------------------------
# Structured parsers (for tabular Excel / CSV data)
# ---------------------------------------------------------------------------

# Column name variations we recognize
COLUMN_ALIASES = {
    "phase": ["phase", "section", "checklist", "checklist phase", "step", "group"],
    "item": ["item", "action", "check", "label", "task", "description", "challenge", "subject"],
    "state": [
        "expectedstate", "expected state", "expected_state", "state", "response",
        "value", "setting", "position", "expectation",
    ],
    "notes": ["notes", "note", "remarks", "comment", "comments", "clue", "details"],
    "category": ["category", "checklist category", "type", "checklist type"],
    "name": ["name", "aircraft", "plane", "airplane"],
    "manufacturer": ["manufacturer", "make", "oem", "brand"],
}


def _find_column(headers: list[str], field: str) -> int:
    """Find the index of a column by checking known aliases."""
    normalized = [h.lower().strip() for h in headers]
    for alias in COLUMN_ALIASES.get(field, [field]):
        if alias in normalized:
            return normalized.index(alias)
    return -1


def _is_structured_data(headers: list[str]) -> bool:
    """Check if headers look like a structured checklist table."""
    phase_col = _find_column(headers, "phase")
    item_col = _find_column(headers, "item")
    return phase_col != -1 and item_col != -1


def parse_structured_rows(
    headers: list[str], rows: list[list[str]]
) -> ParsedChecklist:
    """Parse structured tabular data (from Excel or CSV) into a checklist."""
    col_phase = _find_column(headers, "phase")
    col_item = _find_column(headers, "item")
    col_state = _find_column(headers, "state")
    col_notes = _find_column(headers, "notes")
    col_category = _find_column(headers, "category")
    col_name = _find_column(headers, "name")
    col_mfr = _find_column(headers, "manufacturer")

    checklist = ParsedChecklist()

    phase_map: dict[str, ChecklistPhase] = {}

    for row in rows:
        def get(col: int) -> str:
            if col == -1 or col >= len(row):
                return ""
            return row[col].strip()

        phase_title = get(col_phase)
        item_label = get(col_item)
        if not phase_title or not item_label:
            continue

        # Grab plane metadata from first valid row
        if not checklist.plane_name:
            checklist.plane_name = get(col_name)
            checklist.manufacturer = get(col_mfr)

        if col_category != -1:
            cat = get(col_category)
            if cat and not checklist.category:
                checklist.category = cat

        if phase_title not in phase_map:
            phase_map[phase_title] = ChecklistPhase(phase_title)

        phase_map[phase_title].items.append(
            ChecklistItem(
                label=item_label,
                expected_state=get(col_state),
                notes=get(col_notes),
            )
        )

    checklist.phases = list(phase_map.values())
    checklist.skip_uncategorized = False  # structured source — keep all phases
    return checklist


def parse_csv_text(csv_text: str) -> ParsedChecklist | None:
    """Try to parse CSV text as structured data. Returns None if not structured."""
    reader = csv.reader(io.StringIO(csv_text))
    rows_raw = list(reader)
    if len(rows_raw) < 2:
        return None

    headers = [c.strip() for c in rows_raw[0]]
    if _is_structured_data(headers):
        data_rows = [[c.strip() for c in r] for r in rows_raw[1:] if any(c.strip() for c in r)]
        return parse_structured_rows(headers, data_rows)
    return None


# ---------------------------------------------------------------------------
# Main parse dispatcher
# ---------------------------------------------------------------------------


def parse_file(filepath: str) -> ParsedChecklist:
    """Parse a checklist file and return structured data."""
    path = Path(filepath)
    ext = path.suffix.lower()

    if ext == ".pdf":
        print(f"  Extracting text from PDF: {path.name}")
        text = extract_pdf_text(filepath)
        phases = parse_text_to_checklist(text)
        result = ParsedChecklist()
        result.phases = phases
        # Embed reference tables as a separate variant so they appear under
        # the sub-checklist menu in FlightCheck rather than in the main list.
        table_data = _extract_table_data_as_json(filepath)
        if table_data:
            ref_phase = ChecklistPhase("Reference Tables", category="Reference Tables")
            for label, data_uri in table_data:
                ref_phase.items.append(ChecklistItem(label, notes=data_uri))
            result.phases.append(ref_phase)
            print(f"  Embedded {len(table_data)} reference table(s) as JSON data in 'Reference Tables' variant")
        return result

    elif ext == ".docx":
        print(f"  Extracting text from Word document: {path.name}")
        text = extract_docx_text(filepath)
        phases = parse_text_to_checklist(text)
        result = ParsedChecklist()
        result.phases = phases
        return result

    elif ext in (".xlsx", ".xls"):
        print(f"  Reading Excel file: {path.name}")
        headers, rows = extract_excel_data(filepath)

        # Check if it's structured (has phase/item columns)
        if _is_structured_data(headers):
            print("  Detected structured tabular format")
            return parse_structured_rows(headers, rows)
        else:
            # Fall back to heuristic text parsing
            print("  No structured columns found, using heuristic text parser")
            lines = []
            for row in [headers] + rows:
                lines.append("\t".join(c for c in row if c))
            text = "\n".join(lines)
            phases = parse_text_to_checklist(text)
            result = ParsedChecklist()
            result.phases = phases
            return result

    elif ext == ".csv":
        print(f"  Reading CSV file: {path.name}")
        text = path.read_text(encoding="utf-8-sig")

        # Try structured parse first
        structured = parse_csv_text(text)
        if structured and structured.phases:
            print("  Detected structured CSV format")
            return structured
        else:
            print("  Using heuristic text parser")
            phases = parse_text_to_checklist(text)
            result = ParsedChecklist()
            result.phases = phases
            return result

    elif ext in (".txt", ".text"):
        print(f"  Reading text file: {path.name}")
        text = path.read_text(encoding="utf-8-sig")
        phases = parse_text_to_checklist(text)
        result = ParsedChecklist()
        result.phases = phases
        return result

    elif ext in (".html", ".htm"):
        print(f"  Parsing HTML checklist: {path.name}")
        return parse_html_checklist(filepath)

    else:
        raise ValueError(f"Unsupported file type: {ext}\nSupported: .pdf, .docx, .xlsx, .xls, .csv, .txt, .html")


def parse_google_sheet(url: str) -> ParsedChecklist:
    """Download and parse a Google Sheet."""
    csv_text = download_google_sheet(url)

    # Try structured parse
    structured = parse_csv_text(csv_text)
    if structured and structured.phases:
        print("  Detected structured Google Sheet")
        return structured

    # Fall back to heuristic
    print("  Using heuristic text parser on Google Sheet data")
    phases = parse_text_to_checklist(csv_text)
    result = ParsedChecklist()
    result.phases = phases
    return result


# ---------------------------------------------------------------------------
# Reference table image export
# ---------------------------------------------------------------------------


def _find_table_title(page: object, bbox: tuple, table_num: int, page_num: int) -> str:
    """Extract a title for a table by reading the text immediately above its bounding box.
    Falls back to 'Table N — Page P' if no suitable text is found."""
    try:
        x0, top, x1, bottom = bbox
        # Look at up to 80pt above the table, full page width
        search_top = max(0, top - 80)
        if search_top >= top:
            return f"Table {table_num} — Page {page_num}"
        above = page.crop((0, search_top, page.width, top))
        raw = above.extract_text(x_tolerance=3, y_tolerance=3) or ""
        # Take the last non-empty line (closest to the table)
        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        if lines:
            candidate = lines[-1]
            # Reject lines that look like page numbers or are too short
            if len(candidate) >= 4 and not candidate.isdigit():
                return candidate
    except Exception:
        pass
    return f"Table {table_num} — Page {page_num}"


def _extract_table_data_as_json(filepath: str) -> list[tuple[str, str]]:
    """Extract cell data from PDF tables using pdfplumber and return
    (label, data_uri) pairs encoded as JSON for rendering as styled web
    tables in FlightCheck. No pymupdf required.

    Format: 'data:table/json,[[row0col0, row0col1, ...], ...]'
    """
    if not HAS_PDFPLUMBER:
        return []
    import pdfplumber

    results: list[tuple[str, str]] = []
    with pdfplumber.open(filepath) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            for table_spec in page.find_tables():
                rows = table_spec.extract()
                if not rows:
                    continue
                # Normalise: replace None with "" and strip whitespace
                cleaned = [
                    [cell.strip() if cell else "" for cell in row]
                    for row in rows
                ]
                # Skip degenerate tables (fewer than 2 non-empty cells)
                if sum(1 for row in cleaned for cell in row if cell) < 2:
                    continue
                # Try to find a title by extracting text just above the table
                label = _find_table_title(page, table_spec.bbox, len(results) + 1, page_idx + 1)
                encoded = json.dumps(cleaned, ensure_ascii=False, separators=(",", ":"))
                results.append((label, f"data:table/json,{encoded}"))
    return results


def save_reference_images(filepath: str, output_dir: str) -> list[str]:
    """Find all tables in a PDF (speed/performance reference tables) and save each
    as a JPEG image. Requires pdfplumber (for table detection) and pymupdf (for
    rendering). Returns the list of saved image paths."""
    if not HAS_PDFPLUMBER:
        return []
    if not HAS_FITZ:
        print("  Tip: pip install pymupdf to export reference tables as images")
        return []

    import pdfplumber
    import fitz  # pymupdf

    stem = Path(filepath).stem
    saved: list[str] = []

    with pdfplumber.open(filepath) as pdf:
        doc = fitz.open(filepath)
        for page_idx, page in enumerate(pdf.pages):
            table_specs = page.find_tables()
            if not table_specs:
                continue

            fitz_page = doc[page_idx]
            pw = fitz_page.rect.width
            ph = fitz_page.rect.height

            for ti, table_spec in enumerate(table_specs):
                x0, top, x1, bottom = table_spec.bbox
                pad = 10
                clip = fitz.Rect(
                    max(0, x0 - pad),
                    max(0, top - pad),
                    min(pw, x1 + pad),
                    min(ph, bottom + pad),
                )
                mat = fitz.Matrix(2, 2)  # 2× zoom → ~144 dpi
                pix = fitz_page.get_pixmap(matrix=mat, clip=clip)

                suffix = f"_p{page_idx + 1}_t{ti + 1}" if len(table_specs) > 1 else f"_p{page_idx + 1}"
                path = os.path.join(output_dir, f"{stem}_ref{suffix}.jpg")
                pix.save(path, "jpeg")
                saved.append(path)
                print(f"  Saved reference image: {path}")

        doc.close()

    return saved


# ---------------------------------------------------------------------------
# HTML checklist parser (FlightCheck-style self-contained HTML files)
# ---------------------------------------------------------------------------

from html.parser import HTMLParser as _StdHTMLParser

_EM_DASH = '\u2014'  # —


class _TableExtractor(_StdHTMLParser):
    """Collect all <h3> headings and <table> elements from an HTML fragment in
    document order, associating each table with the nearest preceding <h3>.

    Results are (label, table_data) tuples where table_data is a dict:
      headers     – list[str]       (always present)
      rows        – list[list[str]] (always present)
      rowClasses  – list[str]       (only when at least one "shaded-row" exists)
      preNote     – str             (only when a <div class="small"> precedes)
      postNote    – str             (only when a <div class="small"> follows)
    """

    def __init__(self, default_label: str = "Reference Table"):
        super().__init__(convert_charrefs=True)
        self.default_label = default_label
        self.results: list[tuple[str, dict]] = []

        self._label = default_label
        # ── original table state (variable names unchanged) ──────────────────
        self._table_depth = 0
        self._has_th = False
        self._rows: list[list[str]] = []
        self._row: list[str] = []
        self._cell = ""
        self._in_cell = False
        self._in_h3 = False
        self._h3_buf = ""
        # ── added: per-row shading class ─────────────────────────────────────
        self._cur_row_class = ""           # "shaded-row" or ""
        self._row_classes: list[str] = []  # parallel to self._rows
        # ── added: <div class="small"> note capture ──────────────────────────
        self._in_small_div = False
        self._small_depth = 0   # counts nested divs inside the small div
        self._small_buf = ""
        self._pending_pre_note = ""
        self._after_table = False  # True immediately after a table is saved

    def handle_starttag(self, tag: str, attrs: list) -> None:
        tag = tag.lower()
        attr_class = dict(attrs).get("class", "")

        # ── <div> tracking ───────────────────────────────────────────────────
        if tag == "div":
            if self._in_small_div:
                self._small_depth += 1           # nested div inside note
            elif "small" in attr_class.split() and self._table_depth == 0:
                self._in_small_div = True
                self._small_depth = 1
                self._small_buf = ""
            return  # divs never affect heading/table state

        if self._in_small_div:
            return  # ignore all non-div open-tags inside a note div

        # ── original logic (unchanged) ───────────────────────────────────────
        if tag == "h3" and self._table_depth == 0:
            self._in_h3 = True
            self._h3_buf = ""
        elif tag == "table":
            self._table_depth += 1
            if self._table_depth == 1:
                self._rows = []
                self._row_classes = []
                self._has_th = False
        elif tag == "tr" and self._table_depth == 1:
            self._row = []
            self._cur_row_class = "shaded-row" if "shaded-row" in attr_class.split() else ""
        elif tag in ("th", "td") and self._table_depth == 1:
            self._in_cell = True
            self._cell = ""
            if tag == "th":
                self._has_th = True

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()

        # ── <div> tracking ───────────────────────────────────────────────────
        if tag == "div":
            if self._in_small_div:
                self._small_depth -= 1
                if self._small_depth == 0:
                    self._in_small_div = False
                    note = re.sub(r"\s+", " ", self._small_buf).strip()
                    if note:
                        if self._after_table and self.results:
                            # post-note: attach to the table just saved
                            prev = self.results[-1][1]
                            prev["postNote"] = (
                                (prev.get("postNote") or "") + " " + note
                            ).strip()
                        else:
                            # pre-note: hold until next table is saved
                            self._pending_pre_note = (
                                self._pending_pre_note + " " + note
                            ).strip()
                    self._after_table = False
            return  # divs never affect heading/table state

        if self._in_small_div:
            return  # ignore all non-div close-tags inside a note div

        # ── original logic (unchanged) ───────────────────────────────────────
        if tag == "h3":
            self._in_h3 = False
            label = self._h3_buf.strip()
            if label:
                self._label = label
        elif tag == "table":
            if self._table_depth == 1:
                # Filter empty rows, keeping row_classes in sync
                pairs = [
                    (r, rc)
                    for r, rc in zip(self._rows, self._row_classes)
                    if any(c for c in r)
                ]
                rows = [p[0] for p in pairs]
                row_classes = [p[1] for p in pairs]

                if rows and not self._has_th:
                    num_cols = max((len(r) for r in rows), default=0)
                    header = (
                        ["Specification", "Value"]
                        if num_cols == 2
                        else [f"Col {i + 1}" for i in range(num_cols)]
                    )
                    rows = [header] + rows
                    row_classes = [""] + row_classes

                if len(rows) > 1:  # header + at least one data row
                    table_data: dict = {
                        "headers": rows[0],
                        "rows": rows[1:],
                    }
                    data_rc = row_classes[1:]
                    if any(c for c in data_rc):
                        table_data["rowClasses"] = data_rc
                    if self._pending_pre_note:
                        table_data["preNote"] = self._pending_pre_note
                        self._pending_pre_note = ""
                    self.results.append((self._label, table_data))
                    self._label = self.default_label  # reset for next table
                    self._after_table = True
            self._table_depth -= 1
        elif tag == "tr" and self._table_depth == 1:
            if self._row:
                self._rows.append(self._row)
                self._row_classes.append(self._cur_row_class)
        elif tag in ("th", "td") and self._table_depth == 1:
            self._in_cell = False
            self._row.append(re.sub(r"\s+", " ", self._cell).strip())

    def handle_data(self, data: str) -> None:
        if self._in_h3:
            self._h3_buf += data
        elif self._in_cell:
            self._cell += data
        elif self._in_small_div:
            self._small_buf += data


def _html_bracket_content(text: str, start: int) -> str:
    """Return the content between the opening bracket at `start` and its
    matching closing bracket, respecting string literals."""
    depth = 0
    in_str = False
    str_char = ""
    i = start
    while i < len(text):
        c = text[i]
        if in_str:
            if c == "\\" :
                i += 2
                continue
            if c == str_char:
                in_str = False
        elif c in ('"', "'"):
            in_str = True
            str_char = c
        elif c in ("[", "{"):
            depth += 1
        elif c in ("]", "}"):
            depth -= 1
            if depth == 0:
                return text[start + 1 : i]
        i += 1
    return ""


def _parse_js_phase(obj_text: str) -> dict | None:
    """Parse a single JS phase object: { name: "...", type: "...", items: [...] }"""
    name_m = re.search(r'name:\s*"([^"]+)"', obj_text)
    type_m = re.search(r'type:\s*"([^"]+)"', obj_text)
    if not name_m or not type_m:
        return None

    items_idx = obj_text.find("items:")
    if items_idx == -1:
        return None
    bracket_idx = obj_text.find("[", items_idx)
    if bracket_idx == -1:
        return None

    items_content = _html_bracket_content(obj_text, bracket_idx)
    items = [m.strip() for m in re.findall(r'"((?:[^"\\]|\\.)*)"', items_content) if m.strip()]

    return {"name": name_m.group(1), "type": type_m.group(1), "items": items}


def _parse_html_js_phases(script_text: str) -> list[ChecklistPhase]:
    """Extract the `const phases = [...]` array from JS and return ChecklistPhase list."""
    decl = script_text.find("const phases = [")
    if decl == -1:
        return []
    arr_start = script_text.find("[", decl)
    if arr_start == -1:
        return []

    arr_content = _html_bracket_content(script_text, arr_start)

    phases: list[ChecklistPhase] = []
    obj_start: int | None = None
    obj_depth = 0
    in_str = False
    str_char = ""
    i = 0

    while i < len(arr_content):
        c = arr_content[i]
        if in_str:
            if c == "\\":
                i += 2
                continue
            if c == str_char:
                in_str = False
        elif c in ('"', "'"):
            in_str = True
            str_char = c
        elif c == "{":
            if obj_depth == 0:
                obj_start = i
            obj_depth += 1
        elif c == "}":
            obj_depth -= 1
            if obj_depth == 0 and obj_start is not None:
                obj_text = arr_content[obj_start : i + 1]
                data = _parse_js_phase(obj_text)
                if data:
                    type_val = data["type"]
                    if type_val == "emergency":
                        category, prefix = "Emergency", "Emergency: "
                    elif type_val == "abnormal":
                        category, prefix = "Abnormal", "Abnormal: "
                    else:
                        category, prefix = "", ""

                    title = prefix + data["name"].title()
                    phase = ChecklistPhase(title, category=category)

                    for raw in data["items"]:
                        if f" {_EM_DASH} " in raw:
                            label, state = raw.split(f" {_EM_DASH} ", 1)
                        else:
                            label, state = raw, ""
                        phase.items.append(ChecklistItem(label.strip(), state.strip()))

                    if phase.items:
                        phases.append(phase)
                obj_start = None
        i += 1

    return phases


def _parse_html_ref_tables(script_text: str) -> list[tuple[str, str]]:
    """Extract reference tables from AIRCRAFT_PERFORMANCE / CHARTS JS template
    literals. Returns (label, data_uri) pairs in `data:table/json,` format."""
    results: list[tuple[str, str]] = []

    # Matches: title: "LABEL", html: `HTML_CONTENT`
    # Also handles single-quoted titles as a fallback
    entry_re = re.compile(r'title:\s*["\']([^"\']+)["\'],\s*html:\s*`([\s\S]*?)`')
    matches = list(entry_re.finditer(script_text))

    if not matches:
        print("  Note: No title/html template-literal entries found in script blocks.")
        print("        Reference tables will be skipped.")
        return results

    for m in matches:
        section_title = m.group(1).strip()
        html_content = m.group(2)

        extractor = _TableExtractor(default_label=section_title)
        try:
            extractor.feed(html_content)
        except Exception as e:
            print(f"  Warning: Could not parse HTML for '{section_title}': {e}")
            continue

        for label, table_data in extractor.results:
            has_extras = any(k in table_data for k in ("rowClasses", "preNote", "postNote"))
            if has_extras:
                # Rich format — requires updated FlightCheck TypeScript component
                payload = table_data
            else:
                # Legacy array format [[header,...],[row,...],...]
                # Compatible with all versions of the FlightCheck importer
                payload = [table_data["headers"]] + table_data["rows"]
            encoded = json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
            results.append((label, f"data:table/json,{encoded}"))

    return results


def parse_html_checklist(filepath: str) -> ParsedChecklist:
    """Parse a FlightCheck-style self-contained HTML checklist file.

    Expects:
    - Aircraft name in <title> tag (trailing 'Checklist' stripped)
    - Checklist phases in a `const phases = [...]` JS array inside a <script>
      Each phase: { name: "...", type: "normal|emergency|abnormal", items: ["Label — State", ...] }
    - Optional reference tables in AIRCRAFT_PERFORMANCE / CHARTS JS objects
      stored as HTML template literals with <table> elements
    """
    text = Path(filepath).read_text(encoding="utf-8", errors="replace")

    # Plane name from <title>
    title_m = re.search(r"<title[^>]*>(.*?)</title>", text, re.IGNORECASE | re.DOTALL)
    raw_title = title_m.group(1).strip() if title_m else ""
    plane_name = re.sub(r"\s*checklist\s*$", "", raw_title, flags=re.IGNORECASE).strip()
    if not plane_name:
        plane_name = Path(filepath).stem.replace("_", " ").replace("-", " ").title()

    # Collect all <script> text
    script_blocks = re.findall(r"<script[^>]*>([\s\S]*?)</script>", text, re.IGNORECASE)
    script_text = "\n".join(script_blocks)

    if not script_text:
        raise RuntimeError("No <script> blocks found in this HTML file.")

    # Parse checklist phases
    phases = _parse_html_js_phases(script_text)

    # Parse reference tables
    ref_pairs = _parse_html_ref_tables(script_text)
    if ref_pairs:
        ref_phase = ChecklistPhase("Reference Tables", category="Reference Tables")
        for label, data_uri in ref_pairs:
            ref_phase.items.append(ChecklistItem(label, notes=data_uri))
        phases.append(ref_phase)
        print(f"  Embedded {len(ref_pairs)} reference table(s) from HTML performance data")

    if not phases:
        raise RuntimeError(
            "No checklist data found in this HTML file. "
            "The parser expects a `const phases = [...]` array in a <script> block."
        )

    result = ParsedChecklist(plane_name=plane_name)
    result.phases = phases
    result.skip_uncategorized = False  # all HTML phases are intentional
    return result


# ---------------------------------------------------------------------------
# Output writers
# ---------------------------------------------------------------------------

CSV_HEADERS = ["name", "manufacturer", "type", "image", "checklist category", "item", "expectedState", "notes", "phase"]


def write_csv(checklist: ParsedChecklist, output_path: str) -> None:
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(CSV_HEADERS)
        writer.writerows(checklist.to_csv_rows())
    print(f"  Saved CSV: {output_path}")


def write_json(checklist: ParsedChecklist, output_path: str) -> None:
    data = checklist.to_flightcheck_json()
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"  Saved JSON: {output_path}")


# ---------------------------------------------------------------------------
# Display
# ---------------------------------------------------------------------------


def print_summary(checklist: ParsedChecklist) -> None:
    print("\n" + "=" * 60)
    print("PARSED CHECKLIST SUMMARY")
    print("=" * 60)
    if checklist.plane_name:
        print(f"  Aircraft:     {checklist.plane_name}")
    if checklist.manufacturer:
        print(f"  Manufacturer: {checklist.manufacturer}")
    if checklist.plane_type and checklist.plane_type != "GA":
        print(f"  Type:         {checklist.plane_type}")
    if checklist.category:
        print(f"  Category:     {checklist.category}")
    print(f"  Phases:       {len(checklist.phases)}")
    print(f"  Total items:  {checklist.total_items}")
    print("-" * 60)

    for phase in checklist.phases:
        print(f"\n  [{phase.title}] ({len(phase.items)} items)")
        for item in phase.items:
            state_str = f" -> {item.expected_state}" if item.expected_state else ""
            notes_str = f"  ({item.notes})" if item.notes else ""
            print(f"    - {item.label}{state_str}{notes_str}")

    print("\n" + "=" * 60)


def print_preview(checklist: ParsedChecklist, max_items: int = 5) -> None:
    """Print a compact preview of the parsed result."""
    print(f"\n  Found {len(checklist.phases)} phases, {checklist.total_items} total items:")
    for phase in checklist.phases:
        shown = min(len(phase.items), max_items)
        print(f"    [{phase.title}] - {len(phase.items)} items")
        for item in phase.items[:shown]:
            state_str = f" -> {item.expected_state}" if item.expected_state else ""
            print(f"      - {item.label}{state_str}")
        if len(phase.items) > shown:
            print(f"      ... and {len(phase.items) - shown} more")


# ---------------------------------------------------------------------------
# Interactive CLI
# ---------------------------------------------------------------------------


def prompt_input(message: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    try:
        value = input(f"{message}{suffix}: ").strip()
    except (EOFError, KeyboardInterrupt):
        print("\nAborted.")
        sys.exit(0)
    return value if value else default


def prompt_choice(message: str, choices: list[str], default: str = "") -> str:
    choices_str = "/".join(
        c.upper() if c == default else c for c in choices
    )
    while True:
        value = prompt_input(f"{message} ({choices_str})", default).lower()
        if value in choices:
            return value
        print(f"  Please choose one of: {', '.join(choices)}")


def interactive_mode() -> None:
    print()
    print("=" * 60)
    print("  Aircraft Checklist Converter")
    print("  Converts PDF, Excel, Google Sheets, Word, HTML -> CSV / JSON")
    print("=" * 60)

    if MISSING_DEPS:
        print(f"\n  Note: Some optional dependencies are not installed.")
        print(f"  Install them for full format support:")
        print(f"    pip install {' '.join(MISSING_DEPS)}")

    # Get source
    print()
    source = prompt_input("Source file path or Google Sheets URL")
    if not source:
        print("No source provided. Exiting.")
        return

    # Parse
    is_url = source.startswith("http://") or source.startswith("https://")
    if is_url:
        if "docs.google.com/spreadsheets" not in source:
            print("  Only Google Sheets URLs are supported.")
            print("  For other URLs, download the file first and pass the file path.")
            return
        checklist = parse_google_sheet(source)
    else:
        if not os.path.isfile(source):
            print(f"  File not found: {source}")
            return
        checklist = parse_file(source)

    if not checklist.phases:
        print("\n  No checklist items could be parsed from this file.")
        print("  Tips:")
        print("    - PDF: Works best with text-based PDFs (not scanned images)")
        print("    - Excel: Add 'phase' and 'item' column headers for best results")
        print("    - Try a different format if available")
        return

    # Show preview
    print_preview(checklist)

    # Prompt for aircraft metadata if not found in source
    if not checklist.plane_name:
        checklist.plane_name = prompt_input("\nAircraft name (e.g., Cessna 172)")
    if not checklist.manufacturer:
        checklist.manufacturer = prompt_input("Manufacturer (e.g., Cessna)")

    # Show full summary?
    show_full = prompt_choice("\nShow full parsed checklist?", ["y", "n"], "n")
    if show_full == "y":
        print_summary(checklist)

    # Output format
    fmt = prompt_choice("\nOutput format", ["csv", "json", "both"], "csv")

    # Output path
    if is_url:
        default_stem = re.sub(r"[^a-zA-Z0-9]", "_", checklist.plane_name or "checklist")
    else:
        default_stem = Path(source).stem

    default_dir = os.getcwd()
    output_dir = prompt_input(f"Output directory", default_dir)
    if not os.path.isdir(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    if fmt in ("csv", "both"):
        csv_path = os.path.join(output_dir, f"{default_stem}.csv")
        write_csv(checklist, csv_path)

    if fmt in ("json", "both"):
        json_path = os.path.join(output_dir, f"{default_stem}.json")
        write_json(checklist, json_path)

    # For PDFs, export any reference/performance tables as JPEG images
    if not is_url and Path(source).suffix.lower() == ".pdf":
        save_reference_images(source, output_dir)

    print("\nDone!")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert aircraft checklist files to structured CSV/JSON.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python checklist_converter.py                              # Interactive
  python checklist_converter.py checklist.pdf                # PDF -> CSV
  python checklist_converter.py checklist.xlsx -f json       # Excel -> JSON
  python checklist_converter.py checklist.docx -f both       # Word -> CSV + JSON
  python checklist_converter.py checklist.html -f both       # HTML -> CSV + JSON
  python checklist_converter.py "https://docs.google.com/spreadsheets/d/..." -f csv

Supported inputs:  .pdf, .docx, .xlsx, .xls, .csv, .txt, .html, Google Sheets URL
Supported outputs: .csv (FlightCheck format), .json (FlightCheck format)
        """,
    )

    parser.add_argument(
        "source",
        nargs="?",
        help="Path to checklist file or Google Sheets URL (omit for interactive mode)",
    )
    parser.add_argument(
        "-f", "--format",
        choices=["csv", "json", "both"],
        default="csv",
        help="Output format (default: csv)",
    )
    parser.add_argument(
        "-o", "--output",
        help="Output directory (default: current directory)",
    )
    parser.add_argument(
        "-n", "--name",
        help="Aircraft name (e.g., 'Cessna 172')",
    )
    parser.add_argument(
        "-m", "--manufacturer",
        help="Manufacturer name (e.g., 'Cessna')",
    )
    parser.add_argument(
        "--full",
        action="store_true",
        help="Print full parsed checklist to console",
    )

    args = parser.parse_args()

    # Interactive mode if no source given
    if not args.source:
        interactive_mode()
        return

    source = args.source
    is_url = source.startswith("http://") or source.startswith("https://")

    # Parse
    print(f"\nParsing: {source}")
    if is_url:
        if "docs.google.com/spreadsheets" not in source:
            print("Error: Only Google Sheets URLs are supported.")
            sys.exit(1)
        checklist = parse_google_sheet(source)
    else:
        if not os.path.isfile(source):
            print(f"Error: File not found: {source}")
            sys.exit(1)
        checklist = parse_file(source)

    if not checklist.phases:
        print("Error: No checklist items could be parsed.")
        sys.exit(1)

    # Apply metadata from args
    if args.name:
        checklist.plane_name = args.name
    if args.manufacturer:
        checklist.manufacturer = args.manufacturer

    # Fill in defaults if still missing
    if not checklist.plane_name:
        if is_url:
            checklist.plane_name = "Unknown Aircraft"
        else:
            checklist.plane_name = Path(source).stem.replace("_", " ").replace("-", " ").title()
    if not checklist.manufacturer:
        checklist.manufacturer = "Unknown"

    # Display
    if args.full:
        print_summary(checklist)
    else:
        print_preview(checklist)

    # Output
    output_dir = args.output or os.getcwd()
    if not os.path.isdir(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    if is_url:
        default_stem = re.sub(r"[^a-zA-Z0-9]", "_", checklist.plane_name)
    else:
        default_stem = Path(source).stem

    if args.format in ("csv", "both"):
        write_csv(checklist, os.path.join(output_dir, f"{default_stem}.csv"))
    if args.format in ("json", "both"):
        write_json(checklist, os.path.join(output_dir, f"{default_stem}.json"))

    # For PDFs, export any reference/performance tables as JPEG images
    if not is_url and Path(source).suffix.lower() == ".pdf":
        save_reference_images(source, output_dir)

    print("\nDone!")


if __name__ == "__main__":
    main()
